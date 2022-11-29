from docx import Document
from docx.shared import Inches

guestnames = ['Илья', 'Мади', 'Артем', 'Искандер']
itr = 0

for x in guestnames:
    document = Document()
    document.add_heading('Воспользуйся возможностью стать богатым!', 0)
    document.add_heading('Объявлен результат конкурса!\n', 1)
    p = document.add_paragraph('Здраствуйте, ' + x + '!\n')
    if itr == 0:
        n = 5000000
        p.add_run('Вы заняли первое место в конкурсе ')
        p.add_run('и стали обладателем главного приза в ')
        p.add_run(str(n) + ' тенге!').bold = True
        document.add_picture('first.png', width=Inches(3))
    elif itr == 1:
        n = 2500000
        p.add_run('Вы заняли второе место в конкурсе')
        p.add_run('и получаете поощрительный приз в ')
        p.add_run(str(n) + ' тенге!').bold = True
        document.add_picture('second.png', width=Inches(3))
    elif itr == 2:
        n = 1250000
        p.add_run('Вы заняли третье место в конкурсе')
        p.add_run('и получаете поощрительный приз в ')
        p.add_run(str(n) + ' тенге!').bold = True
        document.add_picture('third.png', width=Inches(3))
    elif itr == 3:
        p.add_run('К сожалению, вы не попали в тройку лучших, ')
        p.add_run('но за участие в конкурсе мы вам дарим')
        p.add_run(' ноутбук Acer Nitro 5 AN515-57!').bold = True
        document.add_picture('lose.png', width=Inches(2))
        document.add_picture('lose_present.png', width=Inches(2))

    document.add_heading('Спасибо за ваше участие и берегите себя!!!', level=1)
    document.add_page_break()
    document.save(str(itr) + '-' + x + '.docx')
    itr += 1
